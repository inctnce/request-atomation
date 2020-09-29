from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt


class Product:
    def __init__(self, name: str, category: str, specs: list, values: list, extra_info: str, price: float):
        self.name = name
        self.category = category
        self.specs = specs
        self.values = values
        self.extra_info = extra_info
        self.price = price


class Demand:
    def __init__(self, products, quantities, costs, infos):
        self.products = products
        self.quantities = quantities
        self.costs = costs
        self.infos = infos


def create_heading(document: Document, heading_info: str):

    titles_list: list = [
        "Управление-Служба информационных технологий",
        "Руководителю  Контрактной службы С.Н. Можайской",
        "02.07.2020 г. № 30-09-404 О проведении закупки поддержки ПО  Yealink Meeting Server",
        ""
    ]

    header_table = document.add_table(rows=2, cols=2)
    i = 0
    k = 0
    j = 0
    while i < 2:
        row = header_table.rows[i].cells
        while k < 2:
            row[k].add_paragraph(titles_list[j])
            row[k].add_paragraph("")
            row[k].add_paragraph("")
            k += 1
            j += 1
        k = 0
        i += 1

    demand_header = document.add_paragraph()
    demand_header.add_run(heading_info).bold = True
    demand_header.alignment = 1

    document.save("request_files/request.docx")


def create_product_table(document: Document, table_name: str, section_name: str, products: list, quantities: list, costs: list):

    p1 = document.add_paragraph()
    p1.add_run(section_name)
    p1.alignment = 1

    products_table = document.add_table(
        rows=len(products)+1, cols=5)
    products_table.style = "Table Grid"

    department = products_table.rows[0].cells
    department[0].text = "№"
    department[1].text = "Наименование"
    department[2].text = "Кол."
    department[3].text = "Цена, руб."
    department[4].text = "Стоимость, руб."

    i = 0
    k = 0
    while i < len(products):
        row = products_table.rows[i+1].cells
        row[0].text = str(i+1)
        row[1].text = str(products[i].name)

        while k < len(products[i].specs):
            row[1].add_paragraph(str(products[i].specs[k]) + "  " +
                                 str(products[i].values[k]))
            k += 1

        row[2].text = str(quantities[i])
        row[3].text = str(products[i].price)
        row[4].text = str(costs[i])
        i += 1
        k = 0

    products_table.columns[0].width = Inches(2)

    document.save("request_files/request.docx")


def create_info(document: Document, labels_list: list, info_list: list):

    detail_list: list = [
        "",
        "(Поставки товара, места оказания услуг)",
        "(Доставки товара, места оказания услуг): ",
        "",
        "(Соответствующее направление образовательной и научной деятельности)",
        "(ФИО, полное наименование должности, контактный телефон, e-mail)",
        "(В случае закупки товаров)\n(ФИО, полное наименование должности, контактный телефон, e-mail)",
    ]

    i = 1
    while i < len(labels_list):
        print(i)
        p = document.add_paragraph(labels_list[i])
        p.add_run(info_list[i])

        if detail_list[i-1] == "":
            i += 1
            continue

        d = document.add_paragraph()
        d.add_run(detail_list[i-1])

        i += 1

    document.save("request_files/request.docx")


def demand_to_docx(product_list, quantity_list, cost_list, info_list):

    demand = Demand(
        products=product_list,
        quantities=quantity_list,
        costs=cost_list,
        infos=info_list
    )

    document = Document()

    label_list: list = [
        "1) Таблица товаров (работ, услуг):  ",
        "2) Суммарная максимальная стоимость заказа: ",
        "3) Сроки:  ",
        "4) Адрес:  ",
        "5) Источник финансирования:  ",
        "6) Виза заместителя начальника ПФУ:  ",
        "7) Контактное лицо по заявке:  ",
        "8) Материально ответственное лицо:  "
    ]

    create_heading(document, demand.infos[0])

    create_product_table(
        document, 
        demand.infos[0], 
        label_list[0], 
        demand.products, 
        demand.quantities, 
        demand.costs)
    create_info(document, label_list, demand.infos)

    document.add_page_break()

    document.save("request_files/request.docx")
